#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
scipilot-writing-skill :: docx_text.py
======================================
读 / 写 Word(.docx)，服务于中文论文润色（Word 是中文期刊的主流载体）。

  - read_docx(path)  -> [(style, text), ...]   保留 Heading 样式信息
  - to_text(path)    -> str                    纯文本（喂给 writing_lint --mode word）
  - write_docx(paragraphs, path)               纯文本段落写回，零 Markdown

依赖 python-docx（`pip install python-docx`）。缺失时给出清晰提示而非 traceback。

用法
----
    python docx_text.py read paper.docx                 # 打印纯文本
    python docx_text.py read paper.docx --to body.txt   # 导出纯文本供机检
"""
from __future__ import annotations

import argparse
import sys

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass


def _require_docx():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError:
        sys.stderr.write(
            "需要 python-docx 才能读写 .docx：\n    pip install python-docx\n"
            "（若只处理 LaTeX / 纯文本则无需安装。）\n")
        raise SystemExit(3)


def read_docx(path: str) -> list:
    docx = _require_docx()
    doc = docx.Document(path)
    out = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        out.append((style, p.text))
    return out


def to_text(path: str) -> str:
    return "\n".join(t for _, t in read_docx(path) if t.strip())


def write_docx(paragraphs, path: str, heading_prefix: str = "Heading") -> str:
    """paragraphs: 可为 [str,...] 或 [(style, text),...]。纯文本写入，不做任何 Markdown 解析。"""
    docx = _require_docx()
    doc = docx.Document()
    for item in paragraphs:
        if isinstance(item, (tuple, list)):
            style, text = item[0], item[1]
        else:
            style, text = "Normal", item
        if style and style.startswith(heading_prefix):
            try:
                level = int(style.replace(heading_prefix, "").strip() or "1")
            except ValueError:
                level = 1
            doc.add_heading(text, level=min(max(level, 0), 9))
        else:
            doc.add_paragraph(text)
    doc.save(path)
    return path


def _main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Word(.docx) 读写工具")
    sub = ap.add_subparsers(dest="cmd", required=True)
    pr = sub.add_parser("read", help="读 .docx 为纯文本")
    pr.add_argument("path")
    pr.add_argument("--to", default=None, help="导出到文本文件（否则打印 stdout）")
    args = ap.parse_args(argv)

    if args.cmd == "read":
        text = to_text(args.path)
        if args.to:
            open(args.to, "w", encoding="utf-8").write(text)
            sys.stderr.write(f"written: {args.to}\n")
        else:
            print(text)
    return 0


if __name__ == "__main__":
    sys.exit(_main())
